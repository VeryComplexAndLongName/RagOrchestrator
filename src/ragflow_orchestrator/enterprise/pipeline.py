from __future__ import annotations

import importlib
import importlib.util
import json
from collections import defaultdict
from collections.abc import Sequence
from datetime import date, datetime
from pathlib import Path
from typing import Any, Protocol, cast
from urllib import request as urllib_request

from ragflow_orchestrator.models import RetrievalResult
from ragflow_orchestrator.orchestrator import RAGOrchestrator

from .adapters import ACLAwareRetrieverAdapter
from .config import EnterpriseLLMConfig, EnterprisePipelineConfig
from .models import (
    EvidenceItem,
    ReviewBundleRequest,
    ReviewBundleResult,
    ReviewTask,
    RiskAssessmentResult,
    RiskEntry,
    TaskOutput,
)
from .tools import EnterpriseToolRegistry, MCPToolProvider, ToolContext


class QueryRetriever(Protocol):
    def retrieve(
        self,
        question: str,
        top_k: int = 5,
        filters: dict[str, object] | None = None,
    ) -> list[RetrievalResult]:
        ...


class EnterprisePipeline:
    """Adaptive enterprise review bundle for document-centric RAG workflows."""

    def __init__(
        self,
        *,
        query_retriever: QueryRetriever,
        config: EnterprisePipelineConfig | None = None,
        mcp_provider: MCPToolProvider | None = None,
        llm_client: object | None = None,
        intent_llm_client: object | None = None,
        answer_llm_client: object | None = None,
    ) -> None:
        self.query_retriever = query_retriever
        self.config = config or EnterprisePipelineConfig()
        self._mcp_provider = mcp_provider
        self._intent_llm_client = intent_llm_client or llm_client
        self._answer_llm_client = answer_llm_client or llm_client
        self._prompt_orchestrator_available: bool | None = None
        self._last_llm_meta: dict[str, Any] = {}

        self.tools = EnterpriseToolRegistry(
            mcp_provider=mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self.intent_tools = EnterpriseToolRegistry(
            mcp_provider=mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self.answer_tools = EnterpriseToolRegistry(
            mcp_provider=mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self._register_default_tools()

    @classmethod
    def from_orchestrator(
        cls,
        orchestrator: RAGOrchestrator,
        *,
        config: EnterprisePipelineConfig | None = None,
        mcp_provider: MCPToolProvider | None = None,
        llm_client: object | None = None,
        intent_llm_client: object | None = None,
        answer_llm_client: object | None = None,
    ) -> EnterprisePipeline:
        retriever = ACLAwareRetrieverAdapter(
            provider=orchestrator.provider,
            embedder=orchestrator.embedder,
        )
        return cls(
            query_retriever=retriever,
            config=config,
            mcp_provider=mcp_provider,
            llm_client=llm_client,
            intent_llm_client=intent_llm_client,
            answer_llm_client=answer_llm_client,
        )

    def set_config(self, config: EnterprisePipelineConfig) -> None:
        self.config = config
        self.tools = EnterpriseToolRegistry(
            mcp_provider=self._mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self.intent_tools = EnterpriseToolRegistry(
            mcp_provider=self._mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self.answer_tools = EnterpriseToolRegistry(
            mcp_provider=self._mcp_provider,
            enable_mcp=self.config.tooling.enable_mcp,
            strict_failures=self.config.tooling.strict_tool_failures,
        )
        self._register_default_tools()

    def review_bundle(self, request_payload: ReviewBundleRequest) -> ReviewBundleResult:
        tasks = self._resolve_tasks(request_payload)
        evidence = self._retrieve_evidence(request_payload)
        grouped = self._group_evidence(evidence)

        task_outputs = [self._execute_task(task, request_payload, grouped) for task in tasks]

        risk_enabled = (
            request_payload.enable_risk_assessment
            if request_payload.enable_risk_assessment is not None
            else self.config.risk.enabled_by_default
        )
        risk_result = self._run_risk_assessment(task_outputs, grouped, enabled=risk_enabled)

        prompt = self._build_prompt(request_payload, tasks, grouped, task_outputs, risk_result)
        llm_answer = self._generate_answer(prompt, task_outputs, request_payload)

        return ReviewBundleResult(
            requested_tasks=tasks,
            executed_tasks=tasks,
            grouped_evidence=grouped,
            task_outputs=task_outputs,
            risk_assessment=risk_result,
            prompt=prompt,
            llm_answer=llm_answer,
        )

    def call_tool(self, name: str, request_id: str, args: dict[str, Any]) -> dict[str, Any]:
        return self.tools.call(name=name, context=ToolContext(request_id=request_id), args=args)

    # ---------------------------------------------------------------------
    # Internal orchestration
    # ---------------------------------------------------------------------
    def _resolve_tasks(self, req: ReviewBundleRequest) -> list[ReviewTask]:
        if req.requested_tasks:
            tasks = list(req.requested_tasks)
        else:
            llm_tasks = self._analyze_query_with_llm(req)
            if llm_tasks:
                tasks = llm_tasks
            else:
                tasks = [ReviewTask(item) for item in self.config.bundle.default_tasks]
                lowered = req.query_text.lower()
                if any(token in lowered for token in ("риск", "risk", "опас", "угроз")):
                    tasks.append(ReviewTask.RISK_ASSESSMENT)
                if any(token in lowered for token in ("ошиб", "дефект", "problem", "issue")):
                    tasks.append(ReviewTask.PROJECT_ERROR_ANALYSIS)
                if any(token in lowered for token in ("противореч", "consisten", "conflict")):
                    tasks.append(ReviewTask.CONSISTENCY_VALIDATION)

        seen: set[ReviewTask] = set()
        unique: list[ReviewTask] = []
        for task in tasks:
            if task in seen:
                continue
            seen.add(task)
            unique.append(task)
        return unique

    def _retrieve_evidence(self, req: ReviewBundleRequest) -> list[EvidenceItem]:
        base_filters = self._compose_filters(req)
        rows = self.query_retriever.retrieve(
            question=req.query_text,
            top_k=self.config.bundle.top_k_per_retrieval,
            filters=base_filters,
        )

        out: list[EvidenceItem] = []
        for row in rows:
            chunk = row.chunk
            metadata_raw = chunk.metadata if isinstance(chunk.metadata, dict) else {}
            metadata = {str(k): str(v) for k, v in metadata_raw.items()}

            if not self._is_document_allowed(req.document_ids, metadata):
                continue
            if not self._is_valid_on_date(req.as_of_date, metadata):
                continue

            group = metadata.get(self.config.bundle.group_by_metadata_key) or "ungrouped"
            out.append(
                EvidenceItem(
                    id=chunk.id,
                    group=group,
                    content=chunk.text,
                    score=row.score,
                    attribution=self._build_attribution(metadata, chunk.id),
                    metadata=metadata,
                )
            )

        if len(out) > self.config.bundle.max_evidence_total:
            out = sorted(out, key=lambda item: item.score or 0.0, reverse=True)[: self.config.bundle.max_evidence_total]
        return out

    def _compose_filters(self, req: ReviewBundleRequest) -> dict[str, object]:
        filters: dict[str, object] = {}
        if req.department_principals:
            filters["acl_principals"] = req.department_principals
        filters.update(req.extra_context)
        return filters

    def _group_evidence(self, evidence: list[EvidenceItem]) -> dict[str, list[EvidenceItem]]:
        grouped: dict[str, list[EvidenceItem]] = defaultdict(list)
        for item in evidence:
            grouped[item.group].append(item)
        for key in grouped:
            grouped[key] = sorted(grouped[key], key=lambda x: x.score or 0.0, reverse=True)[: self.config.bundle.max_evidence_per_group]
        return dict(grouped)

    def _execute_task(
        self,
        task: ReviewTask,
        req: ReviewBundleRequest,
        grouped: dict[str, list[EvidenceItem]],
    ) -> TaskOutput:
        findings: list[str] = []
        evidence_ids: list[str] = []

        total = sum(len(items) for items in grouped.values())
        if total == 0:
            return TaskOutput(
                task=task,
                summary="No evidence available for this task.",
                findings=["No context chunks found after ACL/validity filtering."],
                evidence_ids=[],
            )

        for group_name, items in grouped.items():
            if not items:
                continue
            evidence_ids.extend(item.id for item in items)
            top = items[0]
            if task == ReviewTask.REQUIREMENTS_EXTRACTION:
                findings.append(f"[{group_name}] Candidate requirement sources identified ({len(items)} chunks).")
            elif task == ReviewTask.COMPLIANCE_CHECK:
                findings.append(f"[{group_name}] Compliance signals detected; top evidence score={top.score or 0.0:.3f}.")
            elif task == ReviewTask.CONSISTENCY_VALIDATION:
                findings.append(f"[{group_name}] Consistency check candidates extracted.")
            elif task == ReviewTask.COMPLETENESS_GAP_ANALYSIS:
                findings.append(f"[{group_name}] Potential coverage gaps to review manually.")
            elif task == ReviewTask.AMBIGUITY_PRECISION_REVIEW:
                findings.append(f"[{group_name}] Ambiguous or weakly specified statements may exist.")
            elif task == ReviewTask.PROJECT_ERROR_ANALYSIS:
                findings.append(f"[{group_name}] Potential project mistakes flagged for review.")
            elif task == ReviewTask.CHANGE_IMPACT_DIFF:
                findings.append(f"[{group_name}] Impact analysis requires baseline/version comparison input.")
            elif task == ReviewTask.RISK_ASSESSMENT:
                findings.append(f"[{group_name}] Risk indicators collected for scoring.")

        summary = f"Task {task.value} processed with {len(evidence_ids)} evidence items."
        return TaskOutput(task=task, summary=summary, findings=findings, evidence_ids=sorted(set(evidence_ids)))

    def _run_risk_assessment(
        self,
        task_outputs: list[TaskOutput],
        grouped: dict[str, list[EvidenceItem]],
        *,
        enabled: bool,
    ) -> RiskAssessmentResult:
        if not enabled:
            return RiskAssessmentResult(enabled=False, entries=[])

        risk_signals = 0
        evidence_ids: list[str] = []
        for task in task_outputs:
            if task.task in {
                ReviewTask.COMPLIANCE_CHECK,
                ReviewTask.CONSISTENCY_VALIDATION,
                ReviewTask.PROJECT_ERROR_ANALYSIS,
                ReviewTask.COMPLETENESS_GAP_ANALYSIS,
            }:
                risk_signals += max(1, len(task.findings))
                evidence_ids.extend(task.evidence_ids)

        severity = "low"
        if risk_signals >= 8:
            severity = "high"
        elif risk_signals >= 4:
            severity = "medium"

        likelihood = min(1.0, 0.2 + (risk_signals * 0.08))
        impact = min(1.0, 0.3 + (len(grouped) * 0.1))
        confidence = min(1.0, 0.5 + (len(evidence_ids) * 0.01))

        score = (
            self.config.risk.severity_weight * (1.0 if severity == "high" else (0.7 if severity == "medium" else 0.4))
            + self.config.risk.likelihood_weight * likelihood
            + self.config.risk.impact_weight * impact
        )

        entry = RiskEntry(
            title="Composite document review risk",
            severity=severity,
            likelihood=round(likelihood, 3),
            impact=round(impact, 3),
            confidence=round(confidence, 3),
            priority_score=round(score, 3),
            mitigation="Review top mismatches and update missing/ambiguous requirements before approval.",
            evidence_ids=sorted(set(evidence_ids))[:20],
        )
        return RiskAssessmentResult(enabled=True, entries=[entry])

    def _build_prompt(
        self,
        req: ReviewBundleRequest,
        tasks: list[ReviewTask],
        grouped: dict[str, list[EvidenceItem]],
        task_outputs: list[TaskOutput],
        risk_result: RiskAssessmentResult,
    ) -> str:
        if self.config.bundle.use_prompt_orchestrator and self._check_prompt_orchestrator_available():
            built = self._build_prompt_via_prompt_orchestrator(req=req, grouped=grouped)
            if built:
                return built

        lines: list[str] = []
        lines.append("You are an enterprise document reviewer. Use evidence only.")
        lines.append(f"User request: {req.query_text}")
        lines.append(f"Tasks: {', '.join(task.value for task in tasks)}")
        lines.append("")

        lines.append("Evidence by group:")
        for group_name, items in grouped.items():
            lines.append(f"- Group: {group_name} ({len(items)} items)")
            for item in items:
                score = f"{item.score:.3f}" if item.score is not None else "n/a"
                lines.append(f"  - [{item.id}] score={score} :: {item.attribution}")
                lines.append(f"    {item.content[:400]}")

        lines.append("")
        lines.append("Intermediate findings:")
        for out in task_outputs:
            lines.append(f"- {out.task.value}: {out.summary}")
            for finding in out.findings[:4]:
                lines.append(f"  - {finding}")

        if risk_result.enabled and risk_result.entries:
            lines.append("")
            lines.append("Risk assessment:")
            for risk in risk_result.entries:
                lines.append(
                    f"- {risk.title}: severity={risk.severity}, likelihood={risk.likelihood}, impact={risk.impact}, priority={risk.priority_score}"
                )

        lines.append("")
        lines.append("Return a structured answer with sections: Summary, Findings, Risks, Recommended actions, Citations.")
        return "\n".join(lines)

    def _build_prompt_via_prompt_orchestrator(
        self,
        *,
        req: ReviewBundleRequest,
        grouped: dict[str, list[EvidenceItem]],
    ) -> str | None:
        try:
            prompt_module = importlib.import_module("prompt_orchestrator")
            state_module = importlib.import_module("prompt_orchestrator.context.state")
        except Exception:
            return None

        PromptBuilder = getattr(prompt_module, "PromptBuilder", None)
        PromptConfig = getattr(prompt_module, "PromptConfig", None)
        PromptContextState = getattr(prompt_module, "PromptContextState", None)
        DocChunk = getattr(state_module, "DocChunk", None)
        if not all((PromptBuilder, PromptConfig, PromptContextState, DocChunk)):
            return None

        PromptBuilderCls = cast(type[Any], PromptBuilder)
        PromptConfigCls = cast(type[Any], PromptConfig)
        PromptContextStateCls = cast(type[Any], PromptContextState)
        DocChunkCls = cast(type[Any], DocChunk)

        rag_chunks: list[Any] = []
        for group_name, items in grouped.items():
            for item in items:
                rag_chunks.append(
                    DocChunkCls(
                        id=item.id,
                        content=f"[{group_name}] {item.content}\nSource: {item.attribution}",
                        score=item.score,
                        metadata=item.metadata,
                    )
                )

        cfg = PromptConfigCls(
            system_prompt="You are an enterprise document reviewer.",
            role="Compliance and Risk Analyst",
            task="Use only provided evidence and include citations for every substantial claim.",
            constraints=[
                "Do not invent facts.",
                "Always include citation IDs in findings.",
                "Highlight uncertainty if evidence is insufficient.",
            ],
            output_format="Markdown",
            examples=[],
        )

        state = PromptContextStateCls(session_id=req.session_id, summary=None, recent_messages=[], rag_chunks=rag_chunks)
        builder = PromptBuilderCls()
        sections = builder.build_sections(config=cfg, state=state, user_message=req.query_text, include_headers=False)
        return "\n\n".join([sections["static"], sections["summary"], sections["recent"], sections["rag"]])

    def _generate_answer(self, prompt: str, task_outputs: list[TaskOutput], req: ReviewBundleRequest) -> str:
        cfg = self.config.answer_llm
        provider = cfg.provider
        if provider == "none":
            merged = [f"{item.task.value}: {item.summary}" for item in task_outputs]
            return "\n".join(merged) if merged else "No findings."

        tool_results: list[dict[str, Any]] = []
        if cfg.enable_tools:
            tool_results = self._plan_and_execute_answer_tools(req=req, base_prompt=prompt)
            if tool_results:
                prompt = f"{prompt}\n\nTool execution results:\n{json.dumps(tool_results, ensure_ascii=False, indent=2)}"

        if provider == "custom":
            client = self._answer_llm_client
            if client is None:
                return "Custom LLM provider selected, but no client was supplied."
            text = client.generate(
                prompt=prompt,
                model=cfg.model,
                max_tokens=cfg.max_tokens,
                temperature=cfg.temperature,
            )
            return str(text).strip()

        if provider == "openai":
            return self._generate_openai(prompt, cfg)
        if provider == "ollama":
            return self._generate_ollama(prompt, cfg)

        return "Unsupported LLM provider configuration."

    # ---------------------------------------------------------------------
    # LLM providers
    # ---------------------------------------------------------------------
    def _generate_openai(self, prompt: str, cfg: EnterpriseLLMConfig) -> str:
        self._last_llm_meta = {}
        try:
            from openai import OpenAI  # type: ignore[import-not-found]
        except ImportError:
            return "OpenAI client is not installed."

        http_client = None
        try:
            import httpx  # type: ignore[import-not-found]

            # Avoid inherited proxy settings that can break local/private vLLM endpoints.
            http_client = httpx.Client(timeout=cfg.timeout_seconds, trust_env=False)
        except Exception:
            http_client = None

        client = OpenAI(
            api_key=cfg.openai_api_key,
            base_url=cfg.openai_base_url,
            http_client=http_client,
        )

        try:
            chat = client.chat.completions.create(
                model=cfg.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=cfg.temperature,
                max_tokens=cfg.max_tokens,
            )
            chat_usage = getattr(chat, "usage", None)
            self._last_llm_meta = {
                "endpoint": "chat.completions",
                "prompt_tokens": int(getattr(chat_usage, "prompt_tokens", 0) or 0),
                "completion_tokens": int(getattr(chat_usage, "completion_tokens", 0) or 0),
                "total_tokens": int(getattr(chat_usage, "total_tokens", 0) or 0),
            }
            return str(chat.choices[0].message.content or "").strip()
        except Exception:
            pass

        response = client.responses.create(
            model=cfg.model,
            input=prompt,
            temperature=cfg.temperature,
            max_output_tokens=cfg.max_tokens,
        )
        usage_obj = getattr(response, "usage", None)
        self._last_llm_meta = {
            "endpoint": "responses",
            "prompt_tokens": int(getattr(usage_obj, "prompt_tokens", 0) or 0),
            "completion_tokens": int(getattr(usage_obj, "completion_tokens", 0) or 0),
            "total_tokens": int(getattr(usage_obj, "total_tokens", 0) or 0),
        }
        text = getattr(response, "output_text", None)
        if text:
            return str(text).strip()
        return ""

    def _generate_ollama(self, prompt: str, cfg: EnterpriseLLMConfig) -> str:
        endpoint = f"{cfg.ollama_base_url.rstrip('/')}/api/generate"
        payload = {
            "model": cfg.model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": cfg.temperature,
                "num_predict": cfg.max_tokens,
            },
        }
        req = urllib_request.Request(
            endpoint,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib_request.urlopen(req, timeout=cfg.timeout_seconds) as resp:
                body = resp.read().decode("utf-8")
                parsed = json.loads(body)
                return str(parsed.get("response", "")).strip()
        except Exception as exc:
            return f"Ollama generation failed: {exc}"

    def _analyze_query_with_llm(self, req: ReviewBundleRequest) -> list[ReviewTask] | None:
        cfg = self.config.intent_llm
        if cfg.provider == "none":
            return None

        available = [task.value for task in ReviewTask]
        prompt = (
            "Analyze enterprise review intent and return strict JSON with keys: "
            "tasks (list[str]), tool_calls (list[object]), notes (string). "
            f"Allowed tasks: {available}. "
            "Return JSON only.\n"
            f"User query: {req.query_text}\n"
            f"Extra context: {json.dumps(req.extra_context, ensure_ascii=False)}"
        )

        raw = self._call_llm_text(prompt=prompt, cfg=cfg, custom_client=self._intent_llm_client)
        parsed = self._safe_json_load(raw)
        tool_calls = self._extract_tool_calls(parsed)
        if cfg.enable_tools and tool_calls:
            tool_results = self._execute_tool_calls(
                registry=self.intent_tools,
                request_id=req.session_id,
                tool_calls=tool_calls,
                cfg=cfg,
            )
            retry_prompt = (
                f"{prompt}\n\nTool results:\n{json.dumps(tool_results, ensure_ascii=False, indent=2)}\n"
                "Using tool results, return final strict JSON only with keys tasks, tool_calls, notes."
            )
            raw = self._call_llm_text(prompt=retry_prompt, cfg=cfg, custom_client=self._intent_llm_client)
            parsed = self._safe_json_load(raw)

        if not isinstance(parsed, dict):
            return None

        candidate = parsed.get("tasks")
        if not isinstance(candidate, list):
            return None

        resolved: list[ReviewTask] = []
        for item in candidate:
            try:
                resolved.append(ReviewTask(str(item)))
            except Exception:
                continue
        return resolved or None

    def _plan_and_execute_answer_tools(self, req: ReviewBundleRequest, base_prompt: str) -> list[dict[str, Any]]:
        cfg = self.config.answer_llm
        available = self.answer_tools.names()
        if not available:
            return []

        planner_prompt = (
            "You are planning tool calls for answer generation. Return strict JSON with key tool_calls only. "
            "Each item: {name: str, args: object}. Return JSON only.\n"
            f"Available tools: {available}\n"
            f"User request: {req.query_text}\n"
            f"Prompt summary: {base_prompt[:1200]}"
        )
        raw = self._call_llm_text(prompt=planner_prompt, cfg=cfg, custom_client=self._answer_llm_client)
        parsed = self._safe_json_load(raw)
        tool_calls = self._extract_tool_calls(parsed)
        if not tool_calls:
            return []
        return self._execute_tool_calls(
            registry=self.answer_tools,
            request_id=req.session_id,
            tool_calls=tool_calls,
            cfg=cfg,
        )

    def _call_llm_text(self, *, prompt: str, cfg: EnterpriseLLMConfig, custom_client: object | None) -> str:
        if cfg.provider == "custom":
            if custom_client is None:
                return ""
            text = custom_client.generate(
                prompt=prompt,
                model=cfg.model,
                max_tokens=cfg.max_tokens,
                temperature=cfg.temperature,
            )
            return str(text).strip()
        if cfg.provider == "openai":
            return self._generate_openai(prompt, cfg)
        if cfg.provider == "ollama":
            return self._generate_ollama(prompt, cfg)
        return ""

    @staticmethod
    def _safe_json_load(raw: str) -> dict[str, Any] | list[Any] | None:
        text = (raw or "").strip()
        if not text:
            return None
        try:
            return json.loads(text)
        except Exception:
            pass

        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            try:
                return json.loads(text[start : end + 1])
            except Exception:
                return None
        return None

    @staticmethod
    def _extract_tool_calls(payload: dict[str, Any] | list[Any] | None) -> list[dict[str, Any]]:
        if not isinstance(payload, dict):
            return []
        calls = payload.get("tool_calls")
        if not isinstance(calls, list):
            return []

        out: list[dict[str, Any]] = []
        for call in calls:
            if not isinstance(call, dict):
                continue
            name = call.get("name")
            args = call.get("args")
            if not isinstance(name, str):
                continue
            if not isinstance(args, dict):
                args = {}
            out.append({"name": name, "args": args})
        return out

    def _execute_tool_calls(
        self,
        *,
        registry: EnterpriseToolRegistry,
        request_id: str,
        tool_calls: Sequence[dict[str, Any]],
        cfg: EnterpriseLLMConfig,
    ) -> list[dict[str, Any]]:
        allowed = set(cfg.allowed_tools)
        limited = list(tool_calls)[: max(0, cfg.max_tool_calls)]
        out: list[dict[str, Any]] = []
        for call in limited:
            name = str(call.get("name") or "")
            args = call.get("args")
            if not isinstance(args, dict):
                args = {}
            if allowed and name not in allowed:
                out.append({"ok": False, "tool": name, "error": "Tool is not in allowed_tools"})
                continue
            res = registry.call(name=name, context=ToolContext(request_id=request_id), args=args)
            out.append({"tool": name, "result": res})
        return out

    # ---------------------------------------------------------------------
    # Utility filters and formatting
    # ---------------------------------------------------------------------
    @staticmethod
    def _build_attribution(metadata: dict[str, str], chunk_id: str) -> str:
        doc = metadata.get("document_id") or metadata.get("doc_id") or "unknown-doc"
        version = metadata.get("version_id") or metadata.get("version") or "unknown-version"
        page = metadata.get("page") or metadata.get("page_number") or "?"
        clause = metadata.get("clause_path") or metadata.get("standard_ref") or ""
        valid_to = metadata.get("valid_to") or ""
        parts = [f"doc={doc}", f"version={version}", f"page={page}", f"chunk={chunk_id}"]
        if clause:
            parts.append(f"clause={clause}")
        if valid_to:
            parts.append(f"valid_to={valid_to}")
        return ", ".join(parts)

    @staticmethod
    def _is_document_allowed(document_ids: list[str], metadata: dict[str, str]) -> bool:
        if not document_ids:
            return True
        doc_id = metadata.get("document_id") or metadata.get("doc_id")
        source_id = metadata.get("source_id")
        return bool(doc_id and doc_id in document_ids) or bool(source_id and source_id in document_ids)

    @staticmethod
    def _is_valid_on_date(as_of_date: date | None, metadata: dict[str, str]) -> bool:
        if as_of_date is None:
            return True
        raw = (metadata.get("valid_to") or "").strip()
        if not raw:
            return True

        # Accept common date strings and ISO datetime.
        candidates = [raw]
        if "T" in raw:
            candidates.append(raw.split("T", 1)[0])

        for item in candidates:
            try:
                valid_to = datetime.fromisoformat(item).date()
                return valid_to >= as_of_date
            except Exception:
                continue
        return True

    def _check_prompt_orchestrator_available(self) -> bool:
        if self._prompt_orchestrator_available is not None:
            return self._prompt_orchestrator_available
        try:
            self._prompt_orchestrator_available = importlib.util.find_spec("prompt_orchestrator") is not None
        except Exception:
            self._prompt_orchestrator_available = False
        return self._prompt_orchestrator_available

    def _register_default_tools(self) -> None:
        if not self.config.tooling.enable_tools:
            return

        self.tools.register("retrieve_context", self._tool_retrieve_context)
        self.tools.register("group_evidence", self._tool_group_evidence)
        self.tools.register("score_risks", self._tool_score_risks)
        self.tools.register("build_attribution", self._tool_build_attribution)

        self.intent_tools.register("retrieve_context", self._tool_retrieve_context)
        self.intent_tools.register("group_evidence", self._tool_group_evidence)
        self.intent_tools.register("score_risks", self._tool_score_risks)
        self.intent_tools.register("build_attribution", self._tool_build_attribution)

        self.answer_tools.register("retrieve_context", self._tool_retrieve_context)
        self.answer_tools.register("group_evidence", self._tool_group_evidence)
        self.answer_tools.register("score_risks", self._tool_score_risks)
        self.answer_tools.register("build_attribution", self._tool_build_attribution)
        self.answer_tools.register("create_word_document", self._tool_create_word_document)
        self.answer_tools.register("create_xlsx_table", self._tool_create_xlsx_table)

    # ---------------------------------------------------------------------
    # Default tools
    # ---------------------------------------------------------------------
    def _tool_retrieve_context(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        question = str(args.get("question") or "")
        top_k = int(args.get("top_k") or self.config.bundle.top_k_per_retrieval)
        filters = args.get("filters")
        if not isinstance(filters, dict):
            filters = None

        rows = self.query_retriever.retrieve(question=question, top_k=top_k, filters=filters)
        return {
            "ok": True,
            "request_id": context.request_id,
            "count": len(rows),
            "items": [
                {
                    "id": row.chunk.id,
                    "score": row.score,
                    "content": row.chunk.text,
                    "metadata": row.chunk.metadata,
                }
                for row in rows
            ],
        }

    def _tool_group_evidence(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        items = args.get("items")
        key = str(args.get("key") or self.config.bundle.group_by_metadata_key)
        if not isinstance(items, list):
            return {"ok": False, "request_id": context.request_id, "error": "items must be a list"}

        grouped: dict[str, int] = defaultdict(int)
        for item in items:
            if not isinstance(item, dict):
                continue
            metadata = item.get("metadata")
            if isinstance(metadata, dict):
                group = str(metadata.get(key) or "ungrouped")
            else:
                group = "ungrouped"
            grouped[group] += 1

        return {"ok": True, "request_id": context.request_id, "groups": dict(grouped)}

    def _tool_score_risks(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        severity = str(args.get("severity") or "medium").lower()
        likelihood = float(args.get("likelihood") or 0.5)
        impact = float(args.get("impact") or 0.5)
        severity_value = 1.0 if severity == "high" else (0.7 if severity == "medium" else 0.4)
        score = (
            self.config.risk.severity_weight * severity_value
            + self.config.risk.likelihood_weight * likelihood
            + self.config.risk.impact_weight * impact
        )
        return {
            "ok": True,
            "request_id": context.request_id,
            "priority_score": round(score, 3),
        }

    def _tool_build_attribution(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        metadata = args.get("metadata")
        chunk_id = str(args.get("chunk_id") or "unknown")
        if not isinstance(metadata, dict):
            return {"ok": False, "request_id": context.request_id, "error": "metadata must be a dict"}

        text = self._build_attribution({str(k): str(v) for k, v in metadata.items()}, chunk_id)
        return {"ok": True, "request_id": context.request_id, "attribution": text}

    def _tool_create_word_document(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        filename = str(args.get("filename") or f"review-{context.request_id}.docx")
        title = str(args.get("title") or "Review Report")
        content = str(args.get("content") or "")
        paragraphs_raw = args.get("paragraphs")
        paragraphs: list[str] = []
        if isinstance(paragraphs_raw, list):
            paragraphs = [str(item) for item in paragraphs_raw]

        try:
            from docx import Document  # type: ignore[import-not-found]
        except ImportError:
            return {
                "ok": False,
                "request_id": context.request_id,
                "error": "python-docx is not installed",
            }

        path = Path(filename)
        if not path.is_absolute():
            path = Path.cwd() / path
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(title, level=1)
        if content:
            doc.add_paragraph(content)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        doc.save(str(path))

        return {
            "ok": True,
            "request_id": context.request_id,
            "path": str(path),
            "bytes": path.stat().st_size,
        }

    def _tool_create_xlsx_table(self, context: ToolContext, args: dict[str, Any]) -> dict[str, Any]:
        filename = str(args.get("filename") or f"review-{context.request_id}.xlsx")
        sheet_name = str(args.get("sheet_name") or "Review")
        headers_raw = args.get("headers")
        rows_raw = args.get("rows")

        headers: list[str] = []
        if isinstance(headers_raw, list):
            headers = [str(item) for item in headers_raw]

        rows: list[list[str]] = []
        if isinstance(rows_raw, list):
            for row in rows_raw:
                if isinstance(row, list):
                    rows.append([str(item) for item in row])

        try:
            from openpyxl import Workbook  # type: ignore[import-not-found]
        except ImportError:
            return {
                "ok": False,
                "request_id": context.request_id,
                "error": "openpyxl is not installed",
            }

        path = Path(filename)
        if not path.is_absolute():
            path = Path.cwd() / path
        path.parent.mkdir(parents=True, exist_ok=True)

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = sheet_name[:31] or "Review"

        if headers:
            worksheet.append(headers)
        for row in rows:
            worksheet.append(row)

        workbook.save(str(path))
        return {
            "ok": True,
            "request_id": context.request_id,
            "path": str(path),
            "rows": len(rows),
            "bytes": path.stat().st_size,
        }
